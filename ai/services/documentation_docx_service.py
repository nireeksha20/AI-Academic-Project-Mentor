from io import BytesIO
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Inches


def clean_markdown(text: str) -> str:
    """
    Remove Markdown formatting accidentally returned by the AI.
    """

    if not text:
        return ""

    # Remove bold / italic Markdown
    text = re.sub(r"\*\*(.*?)\*\*", r"\1", text)
    text = re.sub(r"__(.*?)__", r"\1", text)
    text = re.sub(r"(?<!\*)\*(.*?)\*(?!\*)", r"\1", text)
    text = re.sub(r"(?<!_)_(.*?)_(?!_)", r"\1", text)

    # Remove Markdown heading markers
    text = re.sub(r"^#{1,6}\s*", "", text)

    return text.strip()


def add_centered_text(
    document,
    text,
    size=12,
    bold=False,
):
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = paragraph.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.bold = bold

    return paragraph


def create_docx(
    title: str,
    content: str,
    student_name: str = "",
    document_type: str = "",
    institution_name: str = "",
    department_name: str = "",
) -> bytes:

    document = Document()

    # ==========================================================
    # PAGE SETUP
    # ==========================================================

    section = document.sections[0]

    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    # ==========================================================
    # DEFAULT FONT
    # ==========================================================

    styles = document.styles

    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"].font.size = Pt(12)

    styles["Title"].font.name = "Times New Roman"
    styles["Title"].font.size = Pt(20)

    styles["Heading 1"].font.name = "Times New Roman"
    styles["Heading 1"].font.size = Pt(16)
    styles["Heading 1"].font.bold = True

    styles["Heading 2"].font.name = "Times New Roman"
    styles["Heading 2"].font.size = Pt(14)
    styles["Heading 2"].font.bold = True

    styles["Heading 3"].font.name = "Times New Roman"
    styles["Heading 3"].font.size = Pt(13)
    styles["Heading 3"].font.bold = True

    # ==========================================================
    # COVER PAGE
    # ==========================================================

    # Add some vertical spacing
    for _ in range(4):
        document.add_paragraph()

    # Institution — only if actually supplied
    if institution_name:
        add_centered_text(
            document,
            institution_name,
            size=16,
            bold=True,
        )

    # Department — only if actually supplied
    if department_name:
        add_centered_text(
            document,
            department_name,
            size=13,
            bold=False,
        )

    if institution_name or department_name:
        document.add_paragraph()

    # Document type
    if document_type:
        add_centered_text(
            document,
            document_type.upper(),
            size=17,
            bold=True,
        )

    document.add_paragraph()

    # Project title
    add_centered_text(
        document,
        title,
        size=20,
        bold=True,
    )

    document.add_paragraph()
    document.add_paragraph()

    # Student name
    if student_name:
        add_centered_text(
            document,
            "Prepared By",
            size=12,
            bold=True,
        )

        add_centered_text(
            document,
            student_name,
            size=14,
            bold=True,
        )

    # IMPORTANT:
    # No "Academic Project Documentation" text here.
    # No hard-coded college.
    # No unnecessary generic title.

    document.add_page_break()

    # ==========================================================
    # DOCUMENT CONTENT
    # ==========================================================

    lines = content.splitlines()

    for raw_line in lines:

        line = clean_markdown(raw_line)

        if not line:
            continue

        # ------------------------------------------------------
        # Markdown headings
        # ------------------------------------------------------

        if raw_line.startswith("### "):
            paragraph = document.add_heading(
                clean_markdown(raw_line[4:]),
                level=3,
            )

            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            continue

        if raw_line.startswith("## "):
            paragraph = document.add_heading(
                clean_markdown(raw_line[3:]),
                level=2,
            )

            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            continue

        if raw_line.startswith("# "):
            paragraph = document.add_heading(
                clean_markdown(raw_line[2:]),
                level=1,
            )

            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            continue

        # ------------------------------------------------------
        # Bullet list
        # ------------------------------------------------------

        if line.startswith("- "):
            paragraph = document.add_paragraph(
                clean_markdown(line[2:]),
                style="List Bullet",
            )

            continue

        if line.startswith("• "):
            paragraph = document.add_paragraph(
                clean_markdown(line[2:]),
                style="List Bullet",
            )

            continue

        # ------------------------------------------------------
        # Numbered list
        # ------------------------------------------------------

        numbered_match = re.match(
            r"^\d+\.\s+(.*)",
            line,
        )

        if numbered_match:
            paragraph = document.add_paragraph(
                clean_markdown(numbered_match.group(1)),
                style="List Number",
            )

            continue

        # ------------------------------------------------------
        # Normal paragraph
        # ------------------------------------------------------

        paragraph = document.add_paragraph(
            clean_markdown(line)
        )

        paragraph.paragraph_format.space_after = Pt(8)
        paragraph.paragraph_format.line_spacing = 1.15

    # ==========================================================
    # SAVE DOCUMENT
    # ==========================================================

    output = BytesIO()

    document.save(output)

    return output.getvalue()